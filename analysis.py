import pandas as pd
import math
import re
import docx
from collections import OrderedDict
from stop_words import get_stop_words
import numpy as np
import matplotlib.pyplot as plt

all_data = pd.read_csv('file_name.csv')
df = pd.DataFrame(all_data)
pd.set_option('display.max_colwidth', None)
empty_list = []

for element in all_data['Source']:
  if element in empty_list or element == 'Source ':
    continue
  else:
    try:
      math.isnan(float(element))
    except:
      empty_list.append(element)
    else:
      continue

def events(event, dataframe):
  cleaned_list = []
  for index in dataframe.index[dataframe['Source'] == event].tolist():
    quest_clean = str(df.iloc[[index], [0]]).lstrip()
    for r in (("Question", ""), ("Ask Justin a question", ""), ('?',""), ('.',''), ('(', ''), (')', '')):
      quest_clean = quest_clean.replace(*r)
    new_dict = {event: quest_clean}
    cleaned_list.append(new_dict[event])
  return cleaned_list

def stats(sentence_list):
  frequent_words = []
  word_list = get_stop_words('en')
  word_list.extend(['like','can','tell','next','years',"what’s",'one','view','will','-','&'])
  for sentence in sentence_list:
    for word in sentence.split():
      word = word.lower()
      if word in word_list:
        continue
      else:
        frequent_words.append(word.lower())
  frequency = {i : frequent_words.count(i) for i in frequent_words if frequent_words.count(i)>=2}
  return frequency

mydoc = docx.Document()
for event in empty_list:
  mydoc.add_heading(f'Event is: {event}')
  event_list = events(event, df)
  counter = 0
  while counter < len(event_list):
    event_list[counter] = re.sub("\d+", "", event_list[counter]).lstrip().rstrip()
    counter+=1
  statistics = stats(event_list)
  sorted_stats = OrderedDict(sorted(statistics.items(), key=lambda x: x[1]))
  sorted_stats = list(sorted_stats.items())
  if len(sorted_stats) > 5:
    keyword = list(zip(*sorted_stats))[0]
    keyword = keyword[-5:]
    freq = list(zip(*sorted_stats))[1]
    freq = freq[-5:]
  else:
    keyword = list(zip(*sorted_stats))[0]
    freq = list(zip(*sorted_stats))[1]
  x = np.arange(len(keyword))
  plt.bar(x, freq, align='center')
  plt.xticks(x, keyword) 
  plt.ylabel('Frequency')
  curr_png = f'{event}.png'
  plt.savefig(curr_png)
  plt.show()
  mydoc.add_paragraph(f'Top: {sorted_stats[-1]}')
  mydoc.add_paragraph(f'All: {sorted_stats}')
  mydoc.add_picture(curr_png)
  for element in event_list:
    mydoc.add_paragraph(f'- {element}')
  mydoc.add_paragraph('\n')

mydoc.save("analysis.docx")
